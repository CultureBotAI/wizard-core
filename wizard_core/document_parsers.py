"""Document parsers for PDF / DOCX / HTML / YAML inputs.

Ported from `repo-research-writer/scripts/rrwrite_document_parsers.py`.
Provides a uniform `BaseDocumentParser` interface plus concrete parsers for
each format. Optional dependencies (PyPDF2, python-docx, BeautifulSoup4)
load lazily so importing this module is always safe.

Install extras: ``pip install "wizard-core[parsers]"``
"""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import requests
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


class ParsingError(Exception):
    """Raised when document parsing fails."""


class BaseDocumentParser(ABC):
    """Abstract base for format-specific parsers."""

    def __init__(self, source: str):
        self.source = source
        self.raw_text = ""
        self.structured_data: Dict[str, Any] = {}

    @abstractmethod
    def parse(self) -> Dict[str, Any]:
        """Parse the source and return structured data."""

    def get_text(self) -> str:
        return self.raw_text


class PDFParser(BaseDocumentParser):
    """Extract text from PDF files via PyPDF2."""

    def __init__(self, pdf_path: str):
        if not PDF_AVAILABLE:
            raise ParsingError("PyPDF2 not installed. Install with: pip install PyPDF2")
        super().__init__(pdf_path)
        self.pdf_path = Path(pdf_path)
        if not self.pdf_path.exists():
            raise ParsingError(f"PDF file not found: {pdf_path}")

    def parse(self) -> Dict[str, Any]:
        try:
            pages_text = []
            with open(self.pdf_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    pages_text.append(page.extract_text())
            self.raw_text = "\n".join(pages_text)
            self.structured_data = {
                "text": self.raw_text,
                "pages": pages_text,
                "num_pages": len(pages_text),
            }
            return self.structured_data
        except Exception as e:
            raise ParsingError(f"Failed to parse PDF: {e}") from e


class DOCXParser(BaseDocumentParser):
    """Extract text, paragraphs, and tables from .docx via python-docx."""

    def __init__(self, docx_path: str):
        if not DOCX_AVAILABLE:
            raise ParsingError("python-docx not installed. Install with: pip install python-docx")
        super().__init__(docx_path)
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise ParsingError(f"DOCX file not found: {docx_path}")

    def parse(self) -> Dict[str, Any]:
        try:
            doc = Document(self.docx_path)
            paragraphs = [
                {"text": p.text, "style": p.style.name if p.style else None}
                for p in doc.paragraphs
                if p.text.strip()
            ]
            tables = [
                [[cell.text.strip() for cell in row.cells] for row in table.rows]
                for table in doc.tables
            ]
            self.raw_text = "\n".join(p["text"] for p in paragraphs)
            self.structured_data = {
                "text": self.raw_text,
                "paragraphs": paragraphs,
                "tables": tables,
            }
            return self.structured_data
        except Exception as e:
            raise ParsingError(f"Failed to parse DOCX: {e}") from e


class HTMLParser(BaseDocumentParser):
    """Fetch and parse HTML via requests + BeautifulSoup."""

    def __init__(self, url: str, timeout: int = 30):
        if not HTML_AVAILABLE:
            raise ParsingError(
                "BeautifulSoup4 and requests not installed. "
                "Install with: pip install beautifulsoup4 requests lxml"
            )
        super().__init__(url)
        self.url = url
        self.timeout = timeout

    def parse(self) -> Dict[str, Any]:
        try:
            response = requests.get(self.url, timeout=self.timeout)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, "lxml")
            for element in soup(["script", "style", "nav", "footer"]):
                element.decompose()
            self.raw_text = soup.get_text(separator="\n", strip=True)
            self.structured_data = {
                "text": self.raw_text,
                "soup": soup,
                "sections": self._extract_sections(soup),
                "tables": self._extract_tables(soup),
                "url": self.url,
            }
            return self.structured_data
        except requests.RequestException as e:
            raise ParsingError(f"Failed to fetch URL: {e}") from e
        except Exception as e:
            raise ParsingError(f"Failed to parse HTML: {e}") from e

    @staticmethod
    def _extract_sections(soup: Any) -> List[Dict[str, Any]]:
        sections = []
        for heading in soup.find_all(["h1", "h2", "h3", "h4"]):
            section = {
                "level": int(heading.name[1]),
                "title": heading.get_text(strip=True),
                "content": [],
            }
            for sibling in heading.find_next_siblings():
                if sibling.name in ["h1", "h2", "h3", "h4"]:
                    break
                if sibling.get_text(strip=True):
                    section["content"].append(sibling.get_text(strip=True))
            sections.append(section)
        return sections

    @staticmethod
    def _extract_tables(soup: Any) -> List[List[List[str]]]:
        tables = []
        for table in soup.find_all("table"):
            table_data = []
            for row in table.find_all("tr"):
                row_data = [cell.get_text(strip=True) for cell in row.find_all(["td", "th"])]
                if row_data:
                    table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        return tables


class YAMLConverter(BaseDocumentParser):
    """Load and look up a named entry inside a structured YAML file.

    Used by repo-research-writer for ``schemas/journal_guidelines.yaml`` —
    each top-level entry under ``journals:`` is keyed by a short name. The
    converter resolves either by key match or by ``full_name`` substring
    match. Generic enough to reuse for any tool that has a similar
    ``{name: {entries: {...}}}`` shape.
    """

    def __init__(self, yaml_path: str, entry_name: str, container_key: str = "journals", name_field: str = "full_name"):
        super().__init__(yaml_path)
        self.yaml_path = Path(yaml_path)
        self.entry_name = entry_name
        self.container_key = container_key
        self.name_field = name_field
        if not self.yaml_path.exists():
            raise ParsingError(f"YAML file not found: {yaml_path}")

    def parse(self) -> Dict[str, Any]:
        try:
            with open(self.yaml_path) as f:
                data = yaml.safe_load(f)

            entry = None
            if isinstance(data, dict) and self.container_key in data:
                entries = data[self.container_key]
                key = self.entry_name.lower().replace(" ", "_")
                if key in entries:
                    entry = entries[key]
                else:
                    for info in entries.values():
                        if isinstance(info, dict):
                            full = info.get(self.name_field, "")
                            if self.entry_name.lower() in full.lower():
                                entry = info
                                break

            if not entry:
                raise ParsingError(f"Entry '{self.entry_name}' not found in {self.yaml_path}")

            self.structured_data = entry
            self.raw_text = str(entry)
            return self.structured_data
        except yaml.YAMLError as e:
            raise ParsingError(f"Failed to parse YAML: {e}") from e
        except Exception as e:
            raise ParsingError(f"Failed to load YAML: {e}") from e


def create_parser(source: str, source_type: str, **kwargs: Any) -> BaseDocumentParser:
    """Build a parser by source type.

    Args:
        source: Path or URL.
        source_type: One of "pdf", "docx", "html"/"url", "yaml".
        **kwargs: Format-specific kwargs (e.g. ``entry_name`` for yaml).
    """
    source_type = source_type.lower()
    if source_type == "pdf":
        return PDFParser(source)
    if source_type == "docx":
        return DOCXParser(source)
    if source_type in ("html", "url"):
        return HTMLParser(source, **kwargs)
    if source_type == "yaml":
        # Backwards-compatible alias for rrwrite: journal_name -> entry_name
        entry = kwargs.pop("entry_name", None) or kwargs.pop("journal_name", None)
        if not entry:
            raise ParsingError("entry_name (or legacy journal_name) required for YAML parser")
        return YAMLConverter(source, entry, **kwargs)
    raise ParsingError(f"Unknown source type: {source_type}")


# Phase-5 stubs kept for slide-wizard.
def parse_markdown(path: Path) -> Dict[str, Any]:  # pragma: no cover
    raise NotImplementedError("Markdown structural parsing — wire up when needed.")


def parse_pdf(path: Path) -> Dict[str, Any]:
    """Functional wrapper around PDFParser for symmetry with the placeholder API."""
    return PDFParser(str(path)).parse()


def parse_docx(path: Path) -> Dict[str, Any]:
    return DOCXParser(str(path)).parse()


def parse_pptx(path: Path) -> Dict[str, Any]:  # pragma: no cover
    raise NotImplementedError("Implemented in slide-wizard Phase 3 (python-pptx).")
